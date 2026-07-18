"""
╔══════════════════════════════════════════════════════════════════╗
║           بوت الكويز — النسخة المثالية النهائية v3.0           ║
║  Zero-error · AIORateLimiter · Persistence · Full flexibility   ║
╚══════════════════════════════════════════════════════════════════╝

التحسينات على النسخة السابقة:
  ✅ AIORateLimiter — تحكم تلقائي في حدود Telegram بدل retry يدوي
  ✅ PicklePersistence — حفظ البيانات عبر إعادة التشغيل
  ✅ استخراج خيارات عربية (أ. ب. ج. د.)
  ✅ تصحيح bug: normalize_text كان يكسر نهايات الكلمات العربية
  ✅ تصحيح bug: SEND_DELAY global mutation غير آمن
  ✅ تصحيح bug: _send_locks لا تُنظَّف → memory leak
  ✅ دعم (A) بين قوسين كصيغة خيار
  ✅ دعم أسئلة True/False تلقائياً
  ✅ أمر /addadmin لإضافة أدمن بدون إعادة نشر
  ✅ أمر /shuffle لخلط ترتيب الأسئلة
  ✅ شريط تقدم بنسبة مئوية أدق
  ✅ رسالة /help منفصلة وأوضح
  ✅ error handler عالمي مع logging
  ✅ اسم القناة (الخيار الإضافي في كل سؤال) أصبح قابلاً للتغيير من داخل
     البوت مباشرة عبر /label و /setlabel — بدون لمس الكود، لكل شات
     على حدة، ويُحفظ تلقائياً عبر إعادة التشغيل (PicklePersistence)
"""

from __future__ import annotations

import asyncio
import io
import logging
import os
import random
import re
import sys
from dataclasses import dataclass, field
from typing import Any

from telegram import (
    InlineKeyboardButton,
    InlineKeyboardMarkup,
    Update,
)
from telegram.error import BadRequest, Forbidden, RetryAfter, TelegramError
from telegram.ext import (
    AIORateLimiter,
    ApplicationBuilder,
    CallbackQueryHandler,
    CommandHandler,
    ContextTypes,
    MessageHandler,
    PicklePersistence,
    filters,
)

# ── مكتبات اختيارية ────────────────────────────────────────────────────────
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

# ═══════════════════════════════════════════════════════════════════════════
#  السجلات
# ═══════════════════════════════════════════════════════════════════════════
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    handlers=[logging.StreamHandler(sys.stdout)],
)
logger = logging.getLogger("QuizBot")

# ═══════════════════════════════════════════════════════════════════════════
#  الإعدادات
# ═══════════════════════════════════════════════════════════════════════════
TOKEN: str = os.environ["BOT_TOKEN"]


def _parse_ids(*env_names: str) -> set[int]:
    ids: set[int] = set()
    for name in env_names:
        for part in os.environ.get(name, "").split(","):
            part = part.strip()
            if part.lstrip("-").isdigit():
                ids.add(int(part))
    return ids


# مجموعة الأدمن الأساسية من البيئة + الثابت
_BASE_ADMIN_IDS: set[int] = _parse_ids("ADMIN_ID", "ADMIN_IDS") | {8693892771}

DEFAULT_CHANNEL_LABEL = "𝐏𝐬𝐞𝐮𝐝𝐨𝐬𝐜𝐢𝐞𝐧𝐜𝐞"
LABEL_MAX_LEN      = 80  # حد آمن لطول اسم القناة كخيار داخل الاستطلاع
MAIN_CHANNEL       = os.environ.get("MAIN_CHANNEL", "@mj515678")
DEFAULT_SEND_DELAY = float(os.environ.get("SEND_DELAY",        "0.5"))
BUFFER_DELAY       = float(os.environ.get("BUFFER_DELAY",      "2.0"))
IMAGE_BUFFER_DELAY = float(os.environ.get("IMAGE_BUFFER_DELAY","4.0"))
PERSISTENCE_FILE   = os.environ.get("PERSISTENCE_FILE", "bot_data.pkl")

MAX_Q            = 300
MAX_OPT          = 100
TG_MSG_LIMIT     = 4096
TG_MAX_POLL_OPTS = 10
MAX_QUESTIONS    = 500
MAX_SAVED        = 2000

SEND_DESTINATIONS: dict[str, str] = {
    "📡 قناتي الرئيسية": MAIN_CHANNEL,
    "💬 نفس المحادثة":   "SAME_CHAT",
}

# ═══════════════════════════════════════════════════════════════════════════
#  إدارة الأدمن الديناميكية (تُخزن في bot_data للـ persistence)
# ═══════════════════════════════════════════════════════════════════════════
def get_all_admins(bot_data: dict) -> set[int]:
    """يجمع الأدمن الأساسيين مع الديناميكيين."""
    return _BASE_ADMIN_IDS | set(bot_data.get("dynamic_admins", set()))


def is_admin(user_id: int, bot_data: dict | None = None) -> bool:
    if bot_data is None:
        return user_id in _BASE_ADMIN_IDS
    return user_id in get_all_admins(bot_data)


# ═══════════════════════════════════════════════════════════════════════════
#  إدارة "اسم القناة" الديناميكي (الخيار الإضافي في كل سؤال)
# ═══════════════════════════════════════════════════════════════════════════
def get_label_cfg(bot_data: dict, chat_id: int) -> dict:
    """يُعيد إعدادات اسم القناة لهذا الشات (يُنشئها بالقيم الافتراضية إن لم توجد)."""
    cfgs = bot_data.setdefault("channel_label_cfg", {})
    return cfgs.setdefault(chat_id, {"enabled": True, "text": DEFAULT_CHANNEL_LABEL})


def set_label_text(bot_data: dict, chat_id: int, text: str) -> str:
    """يضبط اسم القناة الجديد لهذا الشات ويُفعّله تلقائياً، ويُعيد النص بعد التنظيف."""
    text = text.strip()[:LABEL_MAX_LEN]
    cfg = get_label_cfg(bot_data, chat_id)
    cfg["text"] = text
    cfg["enabled"] = True
    return text


def get_active_label(bot_data: dict, chat_id: int) -> str | None:
    """يُعيد اسم القناة الحالي إن كانت الميزة مُفعّلة، أو None إن كانت مُعطّلة."""
    cfg = get_label_cfg(bot_data, chat_id)
    if not cfg.get("enabled", True):
        return None
    return (cfg.get("text") or DEFAULT_CHANNEL_LABEL)[:LABEL_MAX_LEN]


# ═══════════════════════════════════════════════════════════════════════════
#  بنية السؤال
# ═══════════════════════════════════════════════════════════════════════════
@dataclass
class Question:
    question:    str
    options:     list[str]
    correct:     int
    image:       str | None = None
    explanation: str | None = None

    def is_valid(self) -> bool:
        return (
            bool(self.question.strip())
            and 2 <= len(self.options) <= TG_MAX_POLL_OPTS
            and 0 <= self.correct < len(self.options)
            and all(o.strip() for o in self.options)
        )

    def to_dict(self) -> dict:
        return {
            "question":    self.question,
            "options":     self.options,
            "correct":     self.correct,
            "image":       self.image,
            "explanation": self.explanation,
        }

    @staticmethod
    def from_dict(d: dict) -> "Question":
        return Question(
            question=    d.get("question", ""),
            options=     d.get("options",  []),
            correct=     d.get("correct",  0),
            image=       d.get("image"),
            explanation= d.get("explanation"),
        )


def _ensure_question_objects(qs: list) -> list[Question]:
    result = []
    for q in qs:
        if isinstance(q, Question):
            result.append(q)
        elif isinstance(q, dict):
            result.append(Question.from_dict(q))
        else:
            logger.warning("نوع سؤال غير معروف: %s", type(q))
    return result


# ═══════════════════════════════════════════════════════════════════════════
#  أدوات النص
# ═══════════════════════════════════════════════════════════════════════════
def normalize_text(text: str) -> str:
    """
    تطبيع النص: توحيد أسطر، إزالة BOM وZWS.
    ملاحظة: لا نُوحّد الهمزات/الياء هنا لأن ذلك يُفسد المطابقة
    الدقيقة لنصوص الأسئلة المحتوية على مصطلحات علمية عربية.
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u200b", "").replace("\u200c", "").replace("\ufeff", "")
    return text


def clean_option(opt: str) -> str:
    opt = opt.strip()
    opt = re.sub(r"\s*\.{2,}\s*$", "", opt)
    opt = re.sub(r"\s*\.\s*$",    "", opt)
    opt = re.sub(r"\s+",          " ", opt)
    return opt.strip()


def split_message(text: str, size: int = 4000) -> list[str]:
    if len(text) <= size:
        return [text]
    parts: list[str] = []
    while text:
        if len(text) <= size:
            parts.append(text)
            break
        cut = text.rfind("\n", 0, size)
        if cut < size // 2:
            cut = text.rfind(" ", 0, size)
        if cut < size // 2:
            cut = size
        parts.append(text[:cut])
        text = text[cut:].lstrip("\n ")
    return parts


def extract_text_from_file(filename: str, data: bytes) -> str:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else "txt"

    if ext == "docx":
        if DocxDocument is None:
            raise RuntimeError("python-docx غير مثبتة — أضفها لـ requirements.txt")
        doc = DocxDocument(io.BytesIO(data))
        lines: list[str] = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                lines.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(lines)

    if ext == "pdf":
        if pdfplumber is not None:
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                return "\n".join(p.extract_text() or "" for p in pdf.pages)
        if PdfReader is not None:
            reader = PdfReader(io.BytesIO(data))
            return "\n".join((page.extract_text() or "") for page in reader.pages)
        raise RuntimeError("لا توجد مكتبة PDF — أضف pdfplumber أو PyPDF2 لـ requirements.txt")

    for enc in ("utf-8", "utf-8-sig", "cp1256", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


# ═══════════════════════════════════════════════════════════════════════════
#  تحليل قيمة الإجابة
# ═══════════════════════════════════════════════════════════════════════════
_ALL_PAT = re.compile(
    r"^(?:"
    r"ALL(?:\s+OF\s+(?:THE\s+)?ABOVE)?"
    r"|كل\s*(?:ما\s*)?(?:سبق|ذكر)"
    r"|جميع\s*(?:ما\s*)?(?:سبق|ذكر)"
    r"|كلها|كلهم|جميعها|جميعهم"
    r"|كل\s*الخيارات?\s*(?:صحيحة|صحيح)?"
    r")$",
    re.I | re.U,
)
_NONE_PAT = re.compile(
    r"^(?:"
    r"NONE(?:\s+OF\s+(?:THE\s+)?ABOVE)?"
    r"|لا\s*شيء(?:\s*مما\s*(?:سبق|ذكر))?"
    r"|لا\s*(?:شيء|إجابة)\s*(?:من\s*)?(?:ما\s*)?(?:سبق|صحيحة?)?"
    r")$",
    re.I | re.U,
)
_MULTI_LETTER_PAT = re.compile(
    r"^[A-Ja-j](?:\s*(?:and|or|و|،|,|&|\+|/)\s*[A-Ja-j])+$",
    re.I | re.U,
)
# خريطة الحروف العربية → رقم الخيار
_AR_OPT_MAP = {"أ": 0, "ا": 0, "ب": 1, "ج": 2, "د": 3, "هـ": 4, "ه": 4, "و": 5}

AnswerValue = int | str | None


def parse_answer_value(raw: str) -> AnswerValue:
    s = raw.strip()

    if _ALL_PAT.match(s):
        return "ALL"
    if _NONE_PAT.match(s):
        return "NONE"
    if _MULTI_LETTER_PAT.match(s):
        letters = re.findall(r"\b([A-Ja-j])\b", s, re.I)
        unique  = list(dict.fromkeys(l.upper() for l in letters))
        if len(unique) > 1:
            return f"MULTI:{''.join(unique)}"

    # حرف عربي منفرد
    for ar, idx in _AR_OPT_MAP.items():
        if s == ar:
            return idx

    letters = re.findall(r"\b([A-Ja-j])\b", s, re.I)
    if len(letters) == 1:
        return ord(letters[0].upper()) - ord("A")
    if len(letters) > 1:
        unique = list(dict.fromkeys(l.upper() for l in letters))
        return f"MULTI:{''.join(unique)}"

    return None


def build_multi_option(opts: list[str], letters: str) -> tuple[list[str], int]:
    label = " & ".join(letters)
    combo = f"({label}) كلاهما صحيح ✔"
    new_opts = [clean_option(o) for o in opts] + [combo]
    return new_opts, len(new_opts) - 1


# ═══════════════════════════════════════════════════════════════════════════
#  محلّل الأسئلة
# ═══════════════════════════════════════════════════════════════════════════
# خيارات لاتينية: A) A. A- (A)
_OPT_LAT = re.compile(r"^([A-Ja-j])\s*[.):\-]\s*|^\(([A-Ja-j])\)\s*", re.I)
# خيارات عربية: أ) أ. ب) ب. ج) ج. د)
_OPT_AR  = re.compile(r"^([أابجدهو]|هـ)\s*[.)\-:]\s*", re.U)

_Q_PAT = re.compile(
    r"^(?:Q(?:uestion|s?\.?)?\s*)?(\d+)\s*[.)\-:؟\s]\s*",
    re.I,
)
_ANS_KEYWORD = re.compile(
    r"^(?:"
    r"Correct\s*Answer|Answer|Answers?|Ans|Correct"
    r"|الإجابة\s*الصحيحة?|الإجابة|الجواب|الحل"
    r")\s*[:=\-]\s*",
    re.I | re.U,
)
_ANS_ONLY   = re.compile(r"^([A-Ja-j])\s*$",                  re.I)
_ANS_AR_ONLY= re.compile(r"^([أابجدهو]|هـ)\s*$",              re.U)
_KEY_LINE   = re.compile(r"^(\d+)\s*[-.):\s]\s*([A-Ja-j])\s*$", re.I)
_KEY_HDR    = re.compile(
    r"^(?:answers?\s*key|answer\s*key|key|answers?|الإجابات?|مفتاح\s*الإجابات?)\s*:?\s*$",
    re.I | re.U,
)
_INLINE_OPT = re.compile(r"([A-Ja-j])\s*[.)]\s*(.*?)(?=\s+[A-Ja-j]\s*[.)]|$)", re.I)
_DIVIDER    = re.compile(r"^[-_=*#]{3,}\s*$")
_STEM_START = re.compile(
    r"^(?:Which|What|Who|When|Where|How|Why|Choose|Select|True|False"
    r"|ما|من|أي|اختر|حدد|صح|خطأ|هل)\b",
    re.I | re.U,
)
_TF_PAT     = re.compile(r"^(True|False|صح|خطأ|نعم|لا)\s*$", re.I | re.U)
_EXPL_PAT   = re.compile(
    r"^(?:Explanation|Rationale|Note|ملاحظة|الشرح|التفسير|السبب)\s*[:=\-]\s*",
    re.I | re.U,
)


def _match_option(line: str) -> tuple[bool, str]:
    """يتحقق إن كان السطر خياراً ويُعيد (True, نص_الخيار) أو (False, '')."""
    m = _OPT_LAT.match(line)
    if m:
        return True, line[m.end():].strip()
    m = _OPT_AR.match(line)
    if m:
        return True, line[m.end():].strip()
    return False, ""


def _parse_inline(line: str) -> list[str]:
    m = _INLINE_OPT.findall(line)
    return [clean_option(t) for _, t in m] if len(m) >= 2 else []


def extract_questions(raw_text: str) -> list[Question]:
    text  = normalize_text(raw_text)
    lines = text.splitlines()

    # ── 1. Answer Key منفصل ─────────────────────────────────────────────
    answer_key: dict[int, str] = {}
    key_lines:  set[int]       = set()
    in_key = False

    for i, raw in enumerate(lines):
        s = raw.strip()
        if _KEY_HDR.match(s):
            in_key = True
            key_lines.add(i)
            continue
        if in_key:
            m = _KEY_LINE.match(s)
            if m:
                answer_key[int(m.group(1))] = m.group(2).upper()
                key_lines.add(i)
            elif s and not _DIVIDER.match(s):
                in_key = False

    if len(key_lines) < 3:
        answer_key, key_lines = {}, set()

    # ── 2. تحليل السطر سطراً ────────────────────────────────────────────
    questions: list[Question] = []
    cur_q:     list[str]      = []
    cur_opts:  list[str]      = []
    cur_ans:   AnswerValue    = None
    cur_expl:  str | None     = None
    cur_num:   int | None     = None
    cur_img:   str | None     = None
    opt_idx:   int            = -1
    auto_ctr:  int            = 0

    def flush():
        nonlocal cur_ans
        if not cur_q:
            return
        if cur_ans is None and cur_num is not None and cur_num in answer_key:
            cur_ans = ord(answer_key[cur_num]) - ord("A")
        if cur_ans is None or len(cur_opts) < 2:
            return

        cleaned = [clean_option(o) for o in cur_opts]
        q_text  = " ".join(cur_q).strip()

        def _add(opts: list[str], correct: int):
            if 2 <= len(opts) <= TG_MAX_POLL_OPTS and correct < len(opts):
                questions.append(Question(
                    question=    q_text,
                    options=     opts,
                    correct=     correct,
                    image=       cur_img,
                    explanation= cur_expl,
                ))

        if cur_ans == "ALL":
            _add(cleaned + ["All of the above ✔"], len(cleaned))
        elif cur_ans == "NONE":
            _add(cleaned + ["None of the above ✔"], len(cleaned))
        elif isinstance(cur_ans, str) and cur_ans.startswith("MULTI:"):
            new_opts, correct = build_multi_option(cur_opts, cur_ans[6:])
            _add(new_opts, correct)
        elif isinstance(cur_ans, int):
            _add(cleaned, cur_ans)

    def reset():
        nonlocal cur_q, cur_opts, cur_ans, cur_expl, cur_num, cur_img, opt_idx
        cur_q, cur_opts, cur_ans = [], [], None
        cur_expl, cur_num, cur_img, opt_idx = None, None, None, -1

    for i, raw in enumerate(lines):
        if i in key_lines:
            continue

        line = raw.strip()

        if not line or _DIVIDER.match(line):
            if cur_ans is not None and cur_opts:
                flush(); reset()
            continue

        if _KEY_HDR.match(line):
            if cur_q:
                flush(); reset()
            continue

        # شرح
        m_expl = _EXPL_PAT.match(line)
        if m_expl:
            cur_expl = line[m_expl.end():].strip()
            continue

        # سؤال برقم
        m_q = _Q_PAT.match(line)
        if m_q and not _ANS_KEYWORD.match(line):
            flush()
            auto_ctr += 1
            cur_num = int(m_q.group(1))
            cur_q   = [line[m_q.end():].strip()]
            cur_opts, cur_ans, cur_expl, cur_img, opt_idx = [], None, None, None, -1
            continue

        # إجابة صريحة
        m_ans = _ANS_KEYWORD.match(line)
        if m_ans:
            cur_ans = parse_answer_value(line[m_ans.end():].strip())
            continue

        # خيار (لاتيني أو عربي)
        is_opt, opt_text = _match_option(line)
        if is_opt:
            inline = _parse_inline(line)
            if inline:
                cur_opts = inline
                opt_idx  = len(cur_opts) - 1
            else:
                cur_opts.append(opt_text)
                opt_idx = len(cur_opts) - 1
            continue

        # حرف/حرف عربي وحيد كإجابة بعد الخيارات
        if opt_idx >= 0 and cur_ans is None:
            if _ANS_ONLY.match(line):
                cur_ans = parse_answer_value(line)
                continue
            if _ANS_AR_ONLY.match(line):
                cur_ans = _AR_OPT_MAP.get(line.strip(), None)
                continue

        # True/False تلقائياً → يُنشئ سؤالاً ثنائياً
        if _TF_PAT.match(line) and not cur_opts:
            tf_correct = 0 if re.match(r"^(True|صح|نعم)$", line, re.I | re.U) else 1
            if cur_q:
                cur_opts = ["صح ✅", "خطأ ❌"]
                cur_ans  = tf_correct
            continue

        # استمرار الخيار
        if opt_idx >= 0 and cur_ans is None:
            inline = _parse_inline(line)
            if inline:
                cur_opts = inline
                opt_idx  = len(cur_opts) - 1
            else:
                cur_opts[opt_idx] += " " + line
            continue

        # استمرار السؤال
        if cur_q and cur_ans is None and opt_idx < 0:
            cur_q.append(line)
            continue

        # سؤال جديد بدون رقم
        if _STEM_START.match(line):
            flush()
            auto_ctr += 1
            cur_num = auto_ctr
            cur_q   = [line]
            cur_opts, cur_ans, cur_expl, cur_img, opt_idx = [], None, None, None, -1
            continue

    flush()
    logger.info("✅ استُخرج %d سؤال", len(questions))
    return questions


# ═══════════════════════════════════════════════════════════════════════════
#  محلّل خيارات الصور
# ═══════════════════════════════════════════════════════════════════════════
def extract_options_and_answer(text: str) -> tuple[list[str] | None, int | None]:
    lines   = normalize_text(text).strip().splitlines()
    opts:    list[str] = []
    raw_ans: str | None = None
    opt_idx: int = -1

    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        m_ans = _ANS_KEYWORD.match(line)
        if m_ans:
            raw_ans = line[m_ans.end():].strip()
            continue

        is_opt, opt_text = _match_option(line)
        if is_opt:
            inline = _INLINE_OPT.findall(line)
            if len(inline) >= 2:
                opts    = [clean_option(t) for _, t in inline]
                opt_idx = len(opts) - 1
            else:
                opts.append(clean_option(opt_text))
                opt_idx = len(opts) - 1
            continue

        if opt_idx >= 0 and raw_ans is None and _ANS_ONLY.match(line):
            raw_ans = line
            continue
        if opt_idx >= 0 and raw_ans is None and _ANS_AR_ONLY.match(line):
            raw_ans = line
            continue
        if opt_idx >= 0 and raw_ans is None:
            opts[opt_idx] += " " + line.strip()
            continue

    if len(opts) < 2 or raw_ans is None:
        return None, None

    ans = parse_answer_value(raw_ans)

    if ans == "ALL":
        new = opts + ["All of the above ✔"]
        return new, len(new) - 1
    if ans == "NONE":
        new = opts + ["None of the above ✔"]
        return new, len(new) - 1
    if isinstance(ans, str) and ans.startswith("MULTI:"):
        new, correct = build_multi_option(opts, ans[6:])
        return new, correct
    if isinstance(ans, int) and ans < len(opts):
        return [clean_option(o) for o in opts], ans

    return None, None


# ═══════════════════════════════════════════════════════════════════════════
#  الحماية من التزامن — Lock Pool مع تنظيف تلقائي
# ═══════════════════════════════════════════════════════════════════════════
_send_locks: dict[int, asyncio.Lock] = {}


def get_send_lock(chat_id: int) -> asyncio.Lock:
    if chat_id not in _send_locks:
        _send_locks[chat_id] = asyncio.Lock()
    return _send_locks[chat_id]


def cleanup_lock(chat_id: int):
    """يُزيل Lock بعد الانتهاء لمنع memory leak في حالات الشات الكثيرة."""
    lock = _send_locks.get(chat_id)
    if lock and not lock.locked():
        _send_locks.pop(chat_id, None)


# ═══════════════════════════════════════════════════════════════════════════
#  إرسال الاستطلاعات
# ═══════════════════════════════════════════════════════════════════════════
async def send_polls(
    bot,
    chat_id:         Any,
    questions:       list[Question],
    ctx:             ContextTypes.DEFAULT_TYPE,
    progress_msg=    None,
    start_index:     int = 0,
    control_chat_id: Any = None,
) -> tuple[int, list[dict]]:
    """
    يرسل الأسئلة كـ Quiz Polls.
    - AIORateLimiter يتولى حدود Telegram تلقائياً
    - Exponential back-off للأخطاء العامة
    - شريط تقدم بصري دقيق
    """
    control_chat_id = control_chat_id or chat_id
    flags = ctx.bot_data.setdefault("cancel_flags", {})
    flags[control_chat_id] = False

    # التأخير الخاص بهذا المستخدم أو الافتراضي
    send_delay = ctx.user_data.get("send_delay", DEFAULT_SEND_DELAY)

    failed:  list[dict] = []
    success: int = 0
    total:   int = len(questions)

    for i in range(start_index, total):
        if flags.get(control_chat_id):
            logger.info("إلغاء بواسطة المستخدم عند السؤال %d", i + 1)
            break

        q  = questions[i]
        qn = i + 1

        if not q.is_valid():
            logger.warning("سؤال غير صالح #%d، تخطي", qn)
            failed.append({"index": qn, "question": q.question, "reason": "invalid"})
            continue

        poll_question = q.question
        poll_options  = list(q.options)

        # ── صورة ──────────────────────────────────────────────────────────
        if q.image:
            try:
                cap = (q.question[:1024]) if q.question else "اختر الإجابة الصحيحة 👆"
                await bot.send_photo(chat_id=chat_id, photo=q.image, caption=cap)
            except TelegramError as e:
                logger.warning("فشل إرسال الصورة للسؤال #%d: %s", qn, e)
            poll_question = "اختر الإجابة الصحيحة 👆"

        # ── نص أو خيارات طويلة ───────────────────────────────────────────
        elif len(q.question) > MAX_Q or any(len(o) > MAX_OPT for o in q.options):
            full_text = q.question
            if any(len(o) > MAX_OPT for o in q.options):
                letters   = [chr(ord("A") + k) for k in range(len(q.options))]
                opts_blk  = "\n".join(f"{l}. {o}" for l, o in zip(letters, q.options))
                full_text = f"{q.question}\n\n{opts_blk}"
                poll_options = letters
            for chunk in split_message(full_text):
                try:
                    await bot.send_message(chat_id=chat_id, text=chunk)
                except TelegramError as e:
                    logger.warning("فشل إرسال الرسالة: %s", e)
            if len(q.question) > MAX_Q:
                poll_question = "Choose the correct answer 👆"

        # ── تجهيز الخيارات ───────────────────────────────────────────────
        opts = [o[:MAX_OPT] for o in poll_options]
        channel_label = get_active_label(ctx.bot_data, control_chat_id)
        if channel_label and len(opts) < TG_MAX_POLL_OPTS:
            opts.append(channel_label[:MAX_OPT])

        # ── إرسال الاستطلاع مع retry ─────────────────────────────────────
        sent_ok = False
        retries = 0
        while retries < 5:
            if flags.get(control_chat_id):
                break
            try:
                await bot.send_poll(
                    chat_id           = chat_id,
                    question          = poll_question[:MAX_Q],
                    options           = opts,
                    type              = "quiz",
                    correct_option_id = q.correct,
                    is_anonymous      = True,
                    explanation       = (q.explanation[:200] if q.explanation else None),
                )
                success += 1
                sent_ok = True
                break

            except RetryAfter as e:
                # AIORateLimiter يعالج معظمها، لكن هذا احتياطي
                wait = e.retry_after + 1
                logger.info("RetryAfter %ds للسؤال #%d", wait, qn)
                await asyncio.sleep(wait)
                retries += 1

            except BadRequest as e:
                logger.error("BadRequest للسؤال #%d: %s", qn, e)
                failed.append({"index": qn, "question": q.question, "reason": str(e)})
                break

            except Forbidden as e:
                logger.error("Forbidden — إيقاف الإرسال: %s", e)
                failed.append({"index": qn, "question": q.question, "reason": str(e)})
                flags[control_chat_id] = True
                break

            except TelegramError as e:
                retries += 1
                wait = 2 ** retries
                logger.warning("TelegramError #%d (محاولة %d، انتظار %ds): %s", qn, retries, wait, e)
                await asyncio.sleep(wait)

            except Exception as e:
                logger.exception("خطأ غير متوقع للسؤال #%d", qn)
                failed.append({"index": qn, "question": q.question, "reason": str(e)})
                break

        if sent_ok:
            ctx.user_data["last_sent"] = qn
        elif retries >= 5:
            failed.append({"index": qn, "question": q.question, "reason": "max_retries"})

        # ── شرح الإجابة الطويل كرسالة منفصلة ──────────────────────────────
        if q.explanation and len(q.explanation) > 200:
            try:
                await bot.send_message(
                    chat_id    = chat_id,
                    text       = f"💡 *شرح الإجابة:*\n{q.explanation}",
                    parse_mode = "Markdown",
                )
            except TelegramError:
                pass

        # ── شريط التقدم ──────────────────────────────────────────────────
        if progress_msg and (qn % 10 == 0 or qn == total):
            try:
                pct = int(qn / total * 100)
                filled = pct // 10
                bar = "█" * filled + "░" * (10 - filled)
                await progress_msg.edit_text(
                    f"🚀 [{bar}] {pct}%\n"
                    f"السؤال {qn}/{total} — ✅ {success} نجح | ❌ {len(failed)} فشل"
                )
            except TelegramError:
                pass

        await asyncio.sleep(send_delay)

    flags[control_chat_id] = False
    logger.info("إرسال منتهٍ: ✅ %d | ❌ %d", success, len(failed))
    return success, failed


async def send_failed_file(bot, chat_id: Any, failed: list[dict]):
    if not failed:
        return
    lines = [
        f"#{item['index']}: {item['question']}\n   السبب: {item.get('reason','unknown')}"
        for item in failed
    ]
    buf      = io.BytesIO("\n\n".join(lines).encode("utf-8"))
    buf.name = "failed_questions.txt"
    try:
        await bot.send_document(
            chat_id  = chat_id,
            document = buf,
            filename = "failed_questions.txt",
            caption  = f"📋 {len(failed)} سؤال فشل إرساله",
        )
    except TelegramError as e:
        logger.warning("فشل إرسال ملف الأخطاء: %s", e)


# ═══════════════════════════════════════════════════════════════════════════
#  لوحات المفاتيح
# ═══════════════════════════════════════════════════════════════════════════
def dest_keyboard(prefix: str = "dest") -> InlineKeyboardMarkup:
    rows = [
        [InlineKeyboardButton(label, callback_data=f"{prefix}:{target}")]
        for label, target in SEND_DESTINATIONS.items()
    ]
    if prefix == "dest":
        rows.append([InlineKeyboardButton("📥 حفظ في المحفوظات", callback_data=f"{prefix}:SAVE")])
    rows.append([InlineKeyboardButton("🚫 إلغاء", callback_data=f"{prefix}:CANCEL")])
    return InlineKeyboardMarkup(rows)


# ═══════════════════════════════════════════════════════════════════════════
#  معالجة الصور
# ═══════════════════════════════════════════════════════════════════════════
async def handle_photo(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return

    chat_id = update.effective_chat.id
    photo   = update.message.photo[-1]
    file_id = photo.file_id
    caption = (update.message.caption or "").strip()

    if caption:
        opts, ans = extract_options_and_answer(caption)
        if opts and ans is not None:
            q_text = _extract_question_from_caption(caption)
            await _queue_image_question(update, ctx, file_id, q_text, opts, ans)
            return

    ctx.bot_data.setdefault("pending_images", {})[chat_id] = {
        "file_id": file_id,
        "caption": caption,
    }
    _reset_image_timer(ctx, chat_id)
    await update.message.reply_text(
        "📸 *استلمت الصورة!*\n\n"
        "أرسل الخيارات والإجابة بأي صيغة:\n"
        "```\nA. الخيار الأول\nB. الخيار الثاني\nC. الخيار الثالث\nD. الخيار الرابع\nAnswer: B\n```\n\n"
        "أو بالعربي:\n"
        "```\nأ. الخيار الأول\nب. الخيار الثاني\nج. الخيار الثالث\nالإجابة: ب\n```",
        parse_mode="Markdown",
    )


def _extract_question_from_caption(caption: str) -> str:
    q_lines = []
    for line in caption.splitlines():
        line = line.strip()
        if not line:
            continue
        is_opt, _ = _match_option(line)
        if is_opt or _ANS_KEYWORD.match(line):
            break
        q_lines.append(line)
    return " ".join(q_lines).strip() or "اختر الإجابة الصحيحة"


def _reset_image_timer(ctx: ContextTypes.DEFAULT_TYPE, chat_id: int):
    timers = ctx.bot_data.setdefault("image_timers", {})
    old = timers.get(chat_id)
    if old:
        old.cancel()

    async def _timeout():
        try:
            await asyncio.sleep(IMAGE_BUFFER_DELAY)
        except asyncio.CancelledError:
            return
        if ctx.bot_data.get("pending_images", {}).pop(chat_id, None):
            try:
                await ctx.bot.send_message(
                    chat_id,
                    "⏰ انتهى وقت الانتظار.\n"
                    "أرسل الصورة مجدداً مع الخيارات في الـ caption أو أرسلها ثم الخيارات.",
                )
            except TelegramError:
                pass

    timers[chat_id] = asyncio.create_task(_timeout())


async def _queue_image_question(
    update, ctx, file_id: str, question: str,
    opts: list[str], ans: int,
):
    q = Question(
        question=question or "اختر الإجابة الصحيحة",
        options= opts,
        correct= ans,
        image=   file_id,
    )
    existing: list = ctx.user_data.get("questions", [])
    if len(existing) >= MAX_QUESTIONS:
        await update.message.reply_text(
            f"⚠️ وصلت للحد الأقصى ({MAX_QUESTIONS} سؤال). استخدم /send أو /clear أولاً."
        )
        return

    if not existing:
        ctx.user_data.update({
            "questions":   [q],
            "last_sent":   0,
            "last_failed": [],
            "sending":     False,
            "send_target": None,
        })
        await update.message.reply_text(
            "✅ تم استخراج سؤال مصوّر *(1 سؤال)*\n\n📤 أين تريد الإرسال؟",
            reply_markup=dest_keyboard(),
            parse_mode="Markdown",
        )
    else:
        existing.append(q)
        ctx.user_data["questions"] = existing
        await update.message.reply_text(
            f"✅ أُضيف السؤال المصوّر. الإجمالي: *{len(existing)}*\n"
            "استخدم /send للإرسال أو تابع الإضافة.",
            parse_mode="Markdown",
        )


async def handle_text_after_image(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> bool:
    chat_id = update.effective_chat.id
    pending = ctx.bot_data.get("pending_images", {}).get(chat_id)
    if not pending:
        return False

    text = (update.message.text or "").strip()
    opts, ans = extract_options_and_answer(text)
    if opts is None:
        return False

    timer = ctx.bot_data.get("image_timers", {}).pop(chat_id, None)
    if timer:
        timer.cancel()
    ctx.bot_data["pending_images"].pop(chat_id, None)

    file_id = pending["file_id"]
    q_text  = pending.get("caption", "").strip() or "اختر الإجابة الصحيحة"
    await _queue_image_question(update, ctx, file_id, q_text, opts, ans)
    return True


# ═══════════════════════════════════════════════════════════════════════════
#  أوامر البوت
# ═══════════════════════════════════════════════════════════════════════════
async def cmd_start(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    uid = update.effective_user.id
    if not is_admin(uid, ctx.bot_data):
        await update.message.reply_text("⛔ غير مصرح لك باستخدام هذا البوت.")
        return
    await update.message.reply_text(
        "👋 *أهلاً بك في بوت الكويز!*\n\n"
        "أرسل نصاً أو ملفاً (txt/docx/pdf) أو صورة.\n"
        "للمساعدة الكاملة: /help",
        parse_mode="Markdown",
    )


async def cmd_help(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return
    await update.message.reply_text(
        "📖 *دليل الاستخدام الكامل*\n\n"
        "━━━━━━━━━━━━━━━━━━━━\n"
        "📋 *صيغ الخيارات المدعومة:*\n"
        "`A.` `A)` `A-` `(A)` — لاتينية\n"
        "`أ.` `أ)` `ب.` `ج.` `د.` — عربية\n\n"
        "📋 *صيغ الإجابة:*\n"
        "`Answer: B` · `Ans=C` · `Correct: D`\n"
        "`الإجابة: أ` · حرف/حرف عربي منفرد\n"
        "`Answer: B&D` · `Answer: ALL` · `Answer: NONE`\n\n"
        "📋 *أسئلة True/False:*\n"
        "بعد الإجابة اكتب `True` أو `False` أو `صح` أو `خطأ`\n\n"
        "📋 *الشرح (اختياري):*\n"
        "`Explanation: نص الشرح` أو `الشرح: ...`\n\n"
        "📋 *مفتاح إجابات منفصل:*\n"
        "`Answer Key`\n`1. A`\n`2. B`\n`3. C`\n\n"
        "━━━━━━━━━━━━━━━━━━━━\n"
        "🔧 *الأوامر:*\n"
        "/send — إرسال الأسئلة الحالية\n"
        "/preview — معاينة أول 3 أسئلة\n"
        "/shuffle — خلط ترتيب الأسئلة عشوائياً\n"
        "/saved — عرض المحفوظات\n"
        "/sendsaved — إرسال المحفوظات\n"
        "/clearsaved — مسح المحفوظات\n"
        "/cancel — إيقاف الإرسال فوراً\n"
        "/resume — استئناف آخر عملية\n"
        "/status — الحالة الحالية\n"
        "/stats — إحصائيات الجلسة\n"
        "/clear — مسح الأسئلة الحالية\n"
        "/delay [ث] — ضبط التأخير (0.3–10)\n"
        "/addadmin [ID] — إضافة أدمن مؤقت\n"
        "/removeadmin [ID] — إزالة أدمن مُضاف سابقاً\n"
        "/admins — عرض قائمة كل الأدمن الحاليين\n"
        "/label — إدارة اسم القناة (عرض/تغيير/تعطيل)\n"
        "/setlabel [اسم] — تغيير اسم القناة مباشرة\n"
        "/test — سؤال اختبار\n"
        "/myid — معرفة آيديك\n\n"
        "━━━━━━━━━━━━━━━━━━━━\n"
        "🏷 *اسم القناة (خيار إضافي تلقائي):*\n"
        "كل سؤال يُرسل يحصل تلقائياً على خيار إضافي إجابته غير صحيحة "
        "يحمل اسم قناتك، لأغراض الترويج. غيّره بسهولة:\n"
        "`/setlabel اسم قناتك هنا`\n"
        "أو استخدم /label لقائمة تفاعلية بها تعطيل/تفعيل وإرجاع الافتراضي.",
        parse_mode="Markdown",
    )


async def cmd_myid(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    uid = update.effective_user.id
    await update.message.reply_text(
        f"🆔 آيديك: `{uid}`\n"
        f"{'✅ أنت مشرف' if is_admin(uid, ctx.bot_data) else '⛔ لست مشرفاً'}",
        parse_mode="Markdown",
    )


async def cmd_addadmin(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """إضافة أدمن ديناميكياً (يُحفظ بالـ persistence)."""
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return
    args = ctx.args
    if not args or not args[0].lstrip("-").isdigit():
        await update.message.reply_text(
            "الاستخدام: `/addadmin 123456789`", parse_mode="Markdown"
        )
        return
    new_id = int(args[0])
    admins = ctx.bot_data.setdefault("dynamic_admins", set())
    admins.add(new_id)
    await update.message.reply_text(
        f"✅ تمت إضافة `{new_id}` كأدمن.\n"
        f"إجمالي الأدمن الآن: {len(get_all_admins(ctx.bot_data))}\n\n"
        "لإزالته لاحقاً: `/removeadmin " + str(new_id) + "`",
        parse_mode="Markdown",
    )


async def cmd_removeadmin(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """يزيل أدمناً سبقت إضافته عبر /addadmin (لا يؤثر على الأدمن الأساسيين)."""
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return
    args = ctx.args
    if not args or not args[0].lstrip("-").isdigit():
        await update.message.reply_text(
            "الاستخدام: `/removeadmin 123456789`\n\n"
            "لعرض كل الأدمن الحاليين استخدم /admins",
            parse_mode="Markdown",
        )
        return

    target_id = int(args[0])
    admins    = ctx.bot_data.setdefault("dynamic_admins", set())

    if target_id in _BASE_ADMIN_IDS:
        await update.message.reply_text(
            "⛔ هذا آيدي أدمن أساسي (مضبوط من إعدادات البوت/البيئة) "
            "ولا يمكن إزالته من داخل البوت."
        )
        return

    if target_id not in admins:
        await update.message.reply_text(
            f"⚠️ الآيدي `{target_id}` ليس ضمن الأدمن المُضافين ديناميكياً.\n"
            "استخدم /admins لعرض القائمة الحالية.",
            parse_mode="Markdown",
        )
        return

    admins.discard(target_id)
    await update.message.reply_text(
        f"✅ تمت إزالة `{target_id}` من الأدمن.\n"
        f"إجمالي الأدمن الآن: {len(get_all_admins(ctx.bot_data))}",
        parse_mode="Markdown",
    )


async def cmd_admins(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """يعرض قائمة كل الأدمن الحاليين: الأساسيون (ثابتون) والمُضافون ديناميكياً (قابلون للإزالة)."""
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return
    dynamic = ctx.bot_data.get("dynamic_admins", set())

    lines = ["👥 *قائمة الأدمن الحاليين:*\n", "🔒 *أساسيون (لا يمكن إزالتهم من هنا):*"]
    lines += [f"• `{uid}`" for uid in sorted(_BASE_ADMIN_IDS)] or ["—"]

    lines.append("\n➕ *مُضافون ديناميكياً (قابلون للإزالة):*")
    lines += [f"• `{uid}`" for uid in sorted(dynamic)] if dynamic else ["لا يوجد"]

    lines.append("\nللإضافة: `/addadmin ID`\nللإزالة: `/removeadmin ID`")
    await update.message.reply_text("\n".join(lines), parse_mode="Markdown")


def label_menu_keyboard(enabled: bool) -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("✏️ إرسال اسم جديد", callback_data="label:CHANGE")],
        [InlineKeyboardButton(
            "🚫 تعطيل الخيار الإضافي" if enabled else "✅ تفعيل الخيار الإضافي",
            callback_data="label:TOGGLE",
        )],
        [InlineKeyboardButton("♻️ إرجاع الاسم الافتراضي", callback_data="label:RESET")],
    ])


async def show_label_menu(message, ctx: ContextTypes.DEFAULT_TYPE, chat_id: int):
    cfg    = get_label_cfg(ctx.bot_data, chat_id)
    status = "✅ مُفعّل" if cfg.get("enabled", True) else "🚫 مُعطّل"
    await message.reply_text(
        "🏷 *إدارة اسم القناة*\n\n"
        f"الاسم الحالي: *{cfg.get('text') or DEFAULT_CHANNEL_LABEL}*\n"
        f"الحالة: {status}\n\n"
        "هذا الاسم يُضاف تلقائياً كخيار إضافي (غير صحيح) في نهاية كل سؤال يُرسل — "
        "طريقة سهلة للترويج لقناتك داخل كل استطلاع.\n\n"
        "📌 لتغييره مباشرة اكتب:\n`/setlabel اسمك هنا`\n"
        "أو اضغط الزر أدناه ثم أرسل الاسم الجديد كرسالة عادية.",
        reply_markup=label_menu_keyboard(cfg.get("enabled", True)),
        parse_mode="Markdown",
    )


async def cmd_label(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """يعرض قائمة تفاعلية لإدارة اسم القناة (عرض / تغيير / تعطيل / إرجاع الافتراضي)."""
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return
    await show_label_menu(update.message, ctx, update.effective_chat.id)


async def cmd_setlabel(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """يغيّر اسم القناة مباشرةً: /setlabel اسمك الجديد هنا"""
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return
    chat_id = update.effective_chat.id
    ctx.user_data.pop("awaiting_label", None)

    if not ctx.args:
        await show_label_menu(update.message, ctx, chat_id)
        return

    text = " ".join(ctx.args).strip()
    if not text:
        await update.message.reply_text("⚠️ لا يمكن أن يكون الاسم فارغاً.")
        return

    saved = set_label_text(ctx.bot_data, chat_id, text)
    await update.message.reply_text(
        f"✅ تم تعيين اسم القناة إلى:\n*{saved}*\n\n"
        "سيظهر هذا الآن كخيار إضافي تلقائياً في نهاية كل سؤال يتم إرساله من هذا الشات.",
        parse_mode="Markdown",
    )


async def handle_label_button(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return

    action  = query.data.split(":", 1)[1]
    chat_id = query.message.chat_id
    cfg     = get_label_cfg(ctx.bot_data, chat_id)

    if action == "CHANGE":
        ctx.user_data["awaiting_label"] = True
        await query.message.reply_text(
            "✏️ أرسل الآن الاسم الجديد الذي تريد إضافته كخيار في كل سؤال:\n"
            "_(أرسل /cancel لإلغاء العملية)_",
            parse_mode="Markdown",
        )
        return

    if action == "TOGGLE":
        cfg["enabled"] = not cfg.get("enabled", True)
        try:
            await query.edit_message_text(
                f"{'✅ تم تفعيل' if cfg['enabled'] else '🚫 تم تعطيل'} الخيار الإضافي.\n\n"
                f"الاسم المحفوظ حالياً: *{cfg.get('text') or DEFAULT_CHANNEL_LABEL}*",
                parse_mode="Markdown",
            )
        except TelegramError:
            pass
        return

    if action == "RESET":
        cfg["text"]    = DEFAULT_CHANNEL_LABEL
        cfg["enabled"] = True
        try:
            await query.edit_message_text(
                f"♻️ تم إرجاع الاسم الافتراضي:\n*{DEFAULT_CHANNEL_LABEL}*",
                parse_mode="Markdown",
            )
        except TelegramError:
            pass
        return


async def cmd_cancel(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return
    chat_id = update.effective_chat.id

    for store in ("flush_timers", "image_timers"):
        t = ctx.bot_data.get(store, {}).pop(chat_id, None)
        if t:
            t.cancel()

    ctx.bot_data.get("text_buffers",   {}).pop(chat_id, None)
    ctx.bot_data.get("pending_images", {}).pop(chat_id, None)
    ctx.bot_data.setdefault("cancel_flags", {})[chat_id] = True
    ctx.user_data["sending"] = False
    ctx.user_data.pop("awaiting_label", None)
    cleanup_lock(chat_id)

    await update.message.reply_text(
        "✅ تم إلغاء العملية الحالية.\n"
        "استخدم /resume لاستكمال نفس الدفعة، أو أرسل أسئلة جديدة."
    )


async def cmd_status(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return
    chat_id   = update.effective_chat.id
    questions = _ensure_question_objects(ctx.user_data.get("questions", []))
    saved     = _ensure_question_objects(
        ctx.bot_data.get("saved_questions", {}).get(chat_id, [])
    )
    last_sent = ctx.user_data.get("last_sent", 0)
    failed    = ctx.user_data.get("last_failed", [])
    sending   = ctx.user_data.get("sending", False)
    pending   = chat_id in ctx.bot_data.get("pending_images", {})
    delay     = ctx.user_data.get("send_delay", DEFAULT_SEND_DELAY)

    img_q     = sum(1 for q in questions if q.image)
    saved_img = sum(1 for q in saved if q.image)

    await update.message.reply_text(
        "📊 *الحالة الحالية:*\n\n"
        f"📝 أسئلة جاهزة: *{len(questions)}*\n"
        f"   • نصية: {len(questions) - img_q} | مصوّرة: {img_q}\n"
        f"✅ تم إرساله: {last_sent}\n"
        f"❌ فشل: {len(failed)}\n\n"
        f"📥 *المحفوظات:* {len(saved)} سؤال\n"
        f"   • نصية: {len(saved) - saved_img} | مصوّرة: {saved_img}\n\n"
        f"🚀 إرسال جارٍ: {'نعم ⏳' if sending else 'لا'}\n"
        f"📸 صورة معلّقة: {'نعم ⏳' if pending else 'لا'}\n"
        f"⏱ التأخير: {delay}s",
        parse_mode="Markdown",
    )


async def cmd_stats(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return
    chat_id  = update.effective_chat.id
    session  = ctx.bot_data.get("session_stats", {}).get(chat_id, {})
    total_s  = session.get("total_sent",   0)
    total_f  = session.get("total_failed", 0)
    sessions = session.get("sessions",     0)
    rate     = total_s / max(total_s + total_f, 1) * 100

    await update.message.reply_text(
        "📈 *إحصائيات الجلسة:*\n\n"
        f"📤 إجمالي المُرسل: *{total_s}*\n"
        f"❌ إجمالي الفاشل: *{total_f}*\n"
        f"🔄 عدد عمليات الإرسال: *{sessions}*\n"
        f"📊 معدل النجاح: *{rate:.1f}%*",
        parse_mode="Markdown",
    )


async def cmd_preview(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return
    questions = _ensure_question_objects(ctx.user_data.get("questions", []))
    if not questions:
        await update.message.reply_text("⚠️ لا توجد أسئلة حالياً.")
        return
    n     = min(3, len(questions))
    lines = [f"👁 *معاينة أول {n} أسئلة:*\n"]
    for i, q in enumerate(questions[:n], 1):
        opts_text = "\n".join(
            f"{'✅' if j == q.correct else '  '} {chr(65+j)}. {o}"
            for j, o in enumerate(q.options)
        )
        img_mark = " 🖼" if q.image else ""
        lines.append(f"*{i}.{img_mark}* {q.question[:100]}\n{opts_text}")
        if q.explanation:
            lines.append(f"💡 _{q.explanation[:80]}_")
        lines.append("")
    await update.message.reply_text("\n".join(lines), parse_mode="Markdown")


async def cmd_shuffle(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """خلط ترتيب الأسئلة عشوائياً."""
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return
    if ctx.user_data.get("sending"):
        await update.message.reply_text("⚠️ لا يمكن الخلط أثناء الإرسال.")
        return
    questions = _ensure_question_objects(ctx.user_data.get("questions", []))
    if len(questions) < 2:
        await update.message.reply_text("⚠️ لا يوجد ما يكفي من الأسئلة للخلط.")
        return
    random.shuffle(questions)
    ctx.user_data["questions"] = questions
    await update.message.reply_text(
        f"🔀 تم خلط *{len(questions)}* سؤال عشوائياً.", parse_mode="Markdown"
    )


async def cmd_resume(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return
    if ctx.user_data.get("sending"):
        await update.message.reply_text("⚠️ هناك عملية إرسال جارية بالفعل.")
        return
    questions = _ensure_question_objects(ctx.user_data.get("questions", []))
    last_sent = ctx.user_data.get("last_sent", 0)
    target    = ctx.user_data.get("send_target")
    if not questions or not target:
        await update.message.reply_text("⚠️ لا توجد عملية سابقة لاستئنافها.")
        return
    if last_sent >= len(questions):
        await update.message.reply_text("✅ كل الأسئلة أُرسلت سابقاً.")
        return

    chat_id = update.effective_chat.id
    lock    = get_send_lock(chat_id)
    if lock.locked():
        await update.message.reply_text("⚠️ عملية إرسال جارية.")
        return

    async with lock:
        ctx.bot_data.setdefault("cancel_flags", {})[chat_id] = False
        ctx.user_data["sending"] = True
        msg = await update.message.reply_text(
            f"🔄 استئناف من السؤال *{last_sent + 1}*…", parse_mode="Markdown"
        )
        try:
            success, failed = await send_polls(
                ctx.bot, target, questions, ctx,
                progress_msg=msg, start_index=last_sent, control_chat_id=chat_id,
            )
            ctx.user_data["last_failed"] = failed
            _update_stats(ctx, chat_id, success, len(failed))
            await msg.reply_text(f"✅ تم الإرسال: *{success}*\n❌ فشل: *{len(failed)}*", parse_mode="Markdown")
            await send_failed_file(ctx.bot, chat_id, failed)
        finally:
            ctx.user_data["sending"] = False
            cleanup_lock(chat_id)


async def cmd_send(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return
    if ctx.user_data.get("sending"):
        await update.message.reply_text("⚠️ هناك عملية إرسال جارية.")
        return
    questions = _ensure_question_objects(ctx.user_data.get("questions", []))
    if not questions:
        await update.message.reply_text(
            "⚠️ لا توجد أسئلة حالياً.\nأرسل نصاً أو ملفاً، أو استخدم /sendsaved."
        )
        return
    img_c = sum(1 for q in questions if q.image)
    await update.message.reply_text(
        f"📤 لديك *{len(questions)}* سؤال\n"
        f"_(نصية: {len(questions)-img_c} | مصوّرة: {img_c})_\n\n"
        "📍 أين تريد الإرسال؟",
        reply_markup=dest_keyboard("dest"),
        parse_mode="Markdown",
    )


async def cmd_clear(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return
    if ctx.user_data.get("sending"):
        await update.message.reply_text("⚠️ لا يمكن المسح أثناء الإرسال. استخدم /cancel.")
        return
    count = len(ctx.user_data.get("questions", []))
    if count == 0:
        await update.message.reply_text("📭 القائمة فارغة بالفعل.")
        return
    ctx.user_data.update({
        "questions": [], "last_sent": 0,
        "last_failed": [], "send_target": None,
    })
    await update.message.reply_text(f"🗑 تم مسح *{count}* سؤال.", parse_mode="Markdown")


async def cmd_delay(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return
    args = ctx.args
    if not args:
        current = ctx.user_data.get("send_delay", DEFAULT_SEND_DELAY)
        await update.message.reply_text(
            f"⏱ التأخير الحالي: *{current}s*\n\nلتغييره: `/delay 1.5`\n(المدى: 0.3 – 10 ثانية)",
            parse_mode="Markdown",
        )
        return
    try:
        val = float(args[0])
        if not 0.3 <= val <= 10:
            raise ValueError
        ctx.user_data["send_delay"] = val
        await update.message.reply_text(f"✅ تم ضبط التأخير على *{val}s*", parse_mode="Markdown")
    except ValueError:
        await update.message.reply_text("⚠️ قيمة غير صالحة. مثال: `/delay 1.0`", parse_mode="Markdown")


async def cmd_test(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return
    chat_id = update.effective_chat.id
    label = get_active_label(ctx.bot_data, chat_id) or DEFAULT_CHANNEL_LABEL
    try:
        await ctx.bot.send_poll(
            chat_id           = chat_id,
            question          = "🧪 سؤال اختبار — البوت يعمل بشكل صحيح!",
            options           = ["✅ الإجابة الصحيحة", "❌ خاطئة", "❌ خاطئة", label],
            type              = "quiz",
            correct_option_id = 0,
            is_anonymous      = True,
            explanation       = "هذا مجرد اختبار للتأكد من عمل البوت 🤖",
        )
    except TelegramError as e:
        await update.message.reply_text(f"❌ فشل الاختبار: {e}")


# ═══════════════════════════════════════════════════════════════════════════
#  أوامر المحفوظات
# ═══════════════════════════════════════════════════════════════════════════
async def cmd_saved(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return
    chat_id = update.effective_chat.id
    saved   = _ensure_question_objects(
        ctx.bot_data.get("saved_questions", {}).get(chat_id, [])
    )
    if not saved:
        await update.message.reply_text("📭 المحفوظات فارغة حالياً.")
        return
    img_c   = sum(1 for q in saved if q.image)
    preview = []
    for i, q in enumerate(saved[:5], 1):
        icon = "🖼" if q.image else "📝"
        txt  = q.question[:60] + ("…" if len(q.question) > 60 else "")
        preview.append(f"{i}. {icon} {txt}")
    more = f"\n_…و {len(saved) - 5} سؤال آخر_" if len(saved) > 5 else ""
    await update.message.reply_text(
        f"📥 *المحفوظات:* {len(saved)} سؤال\n"
        f"_(نصية: {len(saved)-img_c} | مصوّرة: {img_c})_\n\n"
        + "\n".join(preview) + more
        + "\n\nاستخدم /sendsaved للإرسال أو /clearsaved للمسح.",
        parse_mode="Markdown",
    )


async def cmd_sendsaved(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return
    if ctx.user_data.get("sending"):
        await update.message.reply_text("⚠️ هناك عملية إرسال جارية.")
        return
    chat_id = update.effective_chat.id
    saved   = _ensure_question_objects(
        ctx.bot_data.get("saved_questions", {}).get(chat_id, [])
    )
    if not saved:
        await update.message.reply_text("📭 المحفوظات فارغة.")
        return
    img_c = sum(1 for q in saved if q.image)
    await update.message.reply_text(
        f"📤 إرسال *{len(saved)}* سؤال من المحفوظات\n"
        f"_(نصية: {len(saved)-img_c} | مصوّرة: {img_c})_\n\n"
        "📍 أين تريد الإرسال؟",
        reply_markup=dest_keyboard("saved_dest"),
        parse_mode="Markdown",
    )


async def cmd_clearsaved(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return
    if ctx.user_data.get("sending"):
        await update.message.reply_text("⚠️ لا يمكن المسح أثناء الإرسال.")
        return
    chat_id = update.effective_chat.id
    saved   = ctx.bot_data.setdefault("saved_questions", {})
    count   = len(saved.get(chat_id, []))
    if count == 0:
        await update.message.reply_text("📭 المحفوظات فارغة.")
        return
    saved[chat_id] = []
    await update.message.reply_text(f"🗑 تم مسح *{count}* سؤال من المحفوظات.", parse_mode="Markdown")


# ═══════════════════════════════════════════════════════════════════════════
#  معالجة النصوص والملفات
# ═══════════════════════════════════════════════════════════════════════════
async def handle_text(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return

    # في انتظار اسم قناة جديد (بعد الضغط على "✏️ إرسال اسم جديد")
    if ctx.user_data.pop("awaiting_label", False):
        chat_id = update.effective_chat.id
        text = (update.message.text or "").strip()
        if not text:
            await update.message.reply_text("⚠️ لا يمكن أن يكون الاسم فارغاً. حاول مجدداً عبر /label")
            return
        saved = set_label_text(ctx.bot_data, chat_id, text)
        await update.message.reply_text(
            f"✅ تم تعيين اسم القناة الجديد:\n*{saved}*\n\n"
            "سيظهر الآن كخيار إضافي تلقائياً في كل سؤال.",
            parse_mode="Markdown",
        )
        return

    if await handle_text_after_image(update, ctx):
        return

    chat_id = update.effective_chat.id
    buffers = ctx.bot_data.setdefault("text_buffers", {})
    timers  = ctx.bot_data.setdefault("flush_timers", {})

    buffers.setdefault(chat_id, []).append(update.message.text or "")

    old = timers.get(chat_id)
    if old:
        old.cancel()

    async def _flush():
        try:
            await asyncio.sleep(BUFFER_DELAY)
        except asyncio.CancelledError:
            return
        chunks = buffers.pop(chat_id, [])
        timers.pop(chat_id, None)
        if chunks:
            await _process(update, ctx, _merge_chunks(chunks))

    timers[chat_id] = asyncio.create_task(_flush())


def _merge_chunks(chunks: list[str]) -> str:
    if not chunks:
        return ""
    merged = chunks[0]
    for i in range(1, len(chunks)):
        prev_full = len(chunks[i - 1]) >= TG_MSG_LIMIT - 100
        merged   += chunks[i] if prev_full else "\n" + chunks[i]
    return merged


async def handle_file(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return

    doc      = update.message.document
    filename = doc.file_name or "file.txt"
    ext      = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    allowed = ("txt", "docx", "pdf", "jpg", "jpeg", "png", "webp")
    if ext not in allowed:
        await update.message.reply_text(
            "⚠️ الصيغ المدعومة: `txt · docx · pdf · jpg · png`",
            parse_mode="Markdown"
        )
        return

    if doc.file_size and doc.file_size > 20 * 1024 * 1024:
        await update.message.reply_text("⚠️ حجم الملف يتجاوز 20MB.")
        return

    # ملفات نصية وdocx — المعالجة العادية بدون AI
    if ext in ("txt", "docx"):
        prog = await update.message.reply_text("⏳ جارٍ قراءة الملف…")
        try:
            tg_file = await doc.get_file()
            data    = bytes(await tg_file.download_as_bytearray())
            text    = extract_text_from_file(filename, data)
        except Exception as e:
            await prog.edit_text(f"⚠️ تعذّر قراءة الملف: {e}")
            return
        await prog.delete()
        await _process(update, ctx, text)
        return

    # PDF وصور — معالجة بالذكاء الاصطناعي
    prog = await update.message.reply_text(
        "🧠 *جارٍ تحليل الملف بالذكاء الاصطناعي…*\n"
        "قد يستغرق 20-40 ثانية ⏳",
        parse_mode="Markdown"
    )
    try:
        tg_file = await doc.get_file()
        data    = bytes(await tg_file.download_as_bytearray())

        from ai_extractor import smart_extract_mcq
        text = await smart_extract_mcq(filename, data)

    except Exception as e:
        logger.exception("خطأ في معالجة الملف بالذكاء الاصطناعي")
        await prog.edit_text(f"⚠️ تعذّر معالجة الملف: {e}")
        return

    await prog.delete()
    await _process(update, ctx, text)

async def _process(update: Update, ctx: ContextTypes.DEFAULT_TYPE, text: str):
    questions = extract_questions(text)
    if not questions:
        await update.message.reply_text(
            "⚠️ لم أجد أسئلة في النص.\n\n"
            "مثال على الصيغة الصحيحة:\n"
            "```\n1. نص السؤال\nA. خيار 1\nB. خيار 2\nC. خيار 3\nAnswer: A\n```",
            parse_mode="Markdown",
        )
        return

    existing = _ensure_question_objects(ctx.user_data.get("questions", []))
    remaining = MAX_QUESTIONS - len(existing)
    if remaining <= 0:
        await update.message.reply_text(
            f"⚠️ وصلت للحد الأقصى ({MAX_QUESTIONS} سؤال). استخدم /send أو /clear."
        )
        return
    if len(questions) > remaining:
        questions = questions[:remaining]
        await update.message.reply_text(
            f"⚠️ تم اقتطاع القائمة إلى {remaining} سؤال (الحد الأقصى {MAX_QUESTIONS})."
        )

    if existing and not ctx.user_data.get("send_target"):
        all_q = existing + questions
        ctx.user_data["questions"] = all_q
        await update.message.reply_text(
            f"✅ تمت إضافة *{len(questions)}* سؤال. الإجمالي: *{len(all_q)}*\n\n"
            "استخدم /send للإرسال أو /preview للمعاينة.",
            parse_mode="Markdown",
        )
        return

    ctx.user_data.update({
        "questions":   questions,
        "last_sent":   0,
        "last_failed": [],
        "sending":     False,
        "send_target": None,
    })
    await update.message.reply_text(
        f"✅ تم استخراج *{len(questions)}* سؤال\n\n📤 أين تريد الإرسال؟",
        reply_markup=dest_keyboard("dest"),
        parse_mode="Markdown",
    )


# ═══════════════════════════════════════════════════════════════════════════
#  معالجة الأزرار
# ═══════════════════════════════════════════════════════════════════════════
async def _do_send(
    bot, chat_id, control_chat_id, questions, ctx,
    progress_msg, label: str,
):
    """دالة مساعدة مشتركة لإرسال الأسئلة."""
    lock = get_send_lock(control_chat_id)
    if lock.locked():
        return False

    q_objs  = _ensure_question_objects(questions)
    img_c   = sum(1 for q in q_objs if q.image)
    ctx.user_data.update({
        "send_target": chat_id,
        "last_sent":   0,
        "sending":     True,
    })
    ctx.bot_data.setdefault("cancel_flags", {})[control_chat_id] = False
    try:
        await progress_msg.edit_text(
            f"🚀 {label} (0/{len(q_objs)})…\n"
            f"📝 نصية: {len(q_objs)-img_c} | 🖼 مصوّرة: {img_c}"
        )
    except TelegramError:
        pass

    async with lock:
        try:
            success, failed = await send_polls(
                bot, chat_id, q_objs, ctx,
                progress_msg=progress_msg, control_chat_id=control_chat_id,
            )
            ctx.user_data["last_failed"] = failed
            _update_stats(ctx, control_chat_id, success, len(failed))
            await progress_msg.reply_text(
                f"✅ تم الإرسال: *{success}*\n❌ فشل: *{len(failed)}*",
                parse_mode="Markdown",
            )
            await send_failed_file(bot, control_chat_id, failed)
            return failed
        finally:
            ctx.user_data["sending"] = False
            cleanup_lock(control_chat_id)


async def handle_destination(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return

    action = query.data.split(":", 1)[1]

    if action == "CANCEL":
        await query.edit_message_text("🚫 تم الإلغاء.")
        return

    if action == "SAVE":
        questions = ctx.user_data.get("questions", [])
        if not questions:
            await query.edit_message_text("⚠️ انتهت الجلسة، أرسل الأسئلة مجدداً.")
            return
        chat_id = query.message.chat_id
        saved   = ctx.bot_data.setdefault("saved_questions", {})
        cur     = saved.get(chat_id, [])
        if len(cur) + len(questions) > MAX_SAVED:
            await query.edit_message_text(
                f"⚠️ المحفوظات ممتلئة ({MAX_SAVED} حد أقصى). استخدم /clearsaved أولاً."
            )
            return
        saved.setdefault(chat_id, []).extend(questions)
        total = len(saved[chat_id])
        ctx.user_data.update({"questions": [], "last_sent": 0, "last_failed": [], "send_target": None})
        await query.edit_message_text(
            f"📥 تم حفظ *{len(questions)}* سؤال! الإجمالي: *{total}*\n\n"
            "استخدم /sendsaved للإرسال لاحقاً.",
            parse_mode="Markdown",
        )
        return

    if ctx.user_data.get("sending"):
        await query.message.reply_text("⚠️ هناك عملية إرسال جارية.")
        return

    questions = ctx.user_data.get("questions", [])
    if not questions:
        await query.edit_message_text("⚠️ انتهت الجلسة، أرسل الأسئلة مجدداً.")
        return

    dest            = query.message.chat_id if action == "SAME_CHAT" else action
    control_chat_id = query.message.chat_id

    await _do_send(ctx.bot, dest, control_chat_id, questions, ctx, query.message, "جاري الإرسال")


async def handle_saved_destination(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    if not is_admin(update.effective_user.id, ctx.bot_data):
        return

    action = query.data.split(":", 1)[1]
    if action == "CANCEL":
        await query.edit_message_text("🚫 تم الإلغاء.")
        return

    if ctx.user_data.get("sending"):
        await query.message.reply_text("⚠️ هناك عملية إرسال جارية.")
        return

    chat_id = query.message.chat_id
    saved   = ctx.bot_data.get("saved_questions", {}).get(chat_id, [])
    if not saved:
        await query.edit_message_text("📭 المحفوظات فارغة.")
        return

    dest = chat_id if action == "SAME_CHAT" else action
    ctx.user_data.update({"questions": list(saved), "last_sent": 0, "last_failed": []})

    failed = await _do_send(ctx.bot, dest, chat_id, saved, ctx, query.message, "جاري إرسال المحفوظات")

    if failed is not None:
        if not failed:
            ctx.bot_data["saved_questions"][chat_id] = []
            await query.message.reply_text("🗑 تم مسح المحفوظات تلقائياً بعد الإرسال الكامل.")
        else:
            failed_idxs = {f["index"] - 1 for f in failed}
            q_objs = _ensure_question_objects(saved)
            ctx.bot_data["saved_questions"][chat_id] = [
                q for idx, q in enumerate(q_objs) if idx in failed_idxs
            ]
            await query.message.reply_text(
                f"⚠️ تم الاحتفاظ بـ {len(failed)} سؤال فاشل في المحفوظات."
            )


# ═══════════════════════════════════════════════════════════════════════════
#  Error Handler العالمي
# ═══════════════════════════════════════════════════════════════════════════
async def error_handler(update: object, ctx: ContextTypes.DEFAULT_TYPE):
    logger.exception("خطأ غير معالَج: %s", ctx.error, exc_info=ctx.error)
    if isinstance(update, Update) and update.effective_message:
        try:
            await update.effective_message.reply_text(
                "⚠️ حدث خطأ غير متوقع. يرجى المحاولة مجدداً أو استخدام /cancel."
            )
        except TelegramError:
            pass


# ═══════════════════════════════════════════════════════════════════════════
#  أدوات داخلية
# ═══════════════════════════════════════════════════════════════════════════
def _update_stats(ctx: ContextTypes.DEFAULT_TYPE, chat_id: int, success: int, failed: int):
    stats = ctx.bot_data.setdefault("session_stats", {})
    s     = stats.setdefault(chat_id, {"total_sent": 0, "total_failed": 0, "sessions": 0})
    s["total_sent"]   += success
    s["total_failed"] += failed
    s["sessions"]     += 1


# ═══════════════════════════════════════════════════════════════════════════
#  main
# ═══════════════════════════════════════════════════════════════════════════
def main():
    persistence = PicklePersistence(filepath=PERSISTENCE_FILE)

    app = (
        ApplicationBuilder()
        .token(TOKEN)
        .persistence(persistence)
        .rate_limiter(AIORateLimiter(max_retries=3))
        .build()
    )

    # ── أوامر ─────────────────────────────────────────────────────────────
    commands = [
        ("start",      cmd_start),
        ("help",       cmd_help),
        ("cancel",     cmd_cancel),
        ("status",     cmd_status),
        ("stats",      cmd_stats),
        ("resume",     cmd_resume),
        ("myid",       cmd_myid),
        ("addadmin",   cmd_addadmin),
        ("removeadmin",cmd_removeadmin),
        ("admins",     cmd_admins),
        ("label",      cmd_label),
        ("setlabel",   cmd_setlabel),
        ("send",       cmd_send),
        ("clear",      cmd_clear),
        ("preview",    cmd_preview),
        ("shuffle",    cmd_shuffle),
        ("saved",      cmd_saved),
        ("sendsaved",  cmd_sendsaved),
        ("clearsaved", cmd_clearsaved),
        ("delay",      cmd_delay),
        ("test",       cmd_test),
    ]
    for cmd, fn in commands:
        app.add_handler(CommandHandler(cmd, fn))

    # ── رسائل ─────────────────────────────────────────────────────────────
    app.add_handler(MessageHandler(filters.PHOTO,                   handle_photo))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))
    app.add_handler(MessageHandler(filters.Document.ALL,            handle_file))

    # ── أزرار ──────────────────────────────────────────────────────────────
    app.add_handler(CallbackQueryHandler(handle_destination,       pattern=r"^dest:"))
    app.add_handler(CallbackQueryHandler(handle_saved_destination, pattern=r"^saved_dest:"))
    app.add_handler(CallbackQueryHandler(handle_label_button,      pattern=r"^label:"))

    # ── Error Handler ──────────────────────────────────────────────────────
    app.add_error_handler(error_handler)

    logger.info("✅ البوت يعمل — AIORateLimiter + PicklePersistence نشطان")
    app.run_polling(drop_pending_updates=True)


if __name__ == "__main__":
    main()
